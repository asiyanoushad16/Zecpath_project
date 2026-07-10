from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import IsAuthenticated
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework.generics import ListAPIView
from rest_framework.pagination import PageNumberPagination
from rest_framework.filters import SearchFilter
from django_filters.rest_framework import DjangoFilterBackend
from .models import User
from .serializers import UserSerializer
from django.db.models import Count
import pdfplumber
import docx
import re
from rest_framework.exceptions import PermissionDenied
from django.core.mail import send_mail
from django.conf import settings
from io import BytesIO


from .models import (
    User,
    Employer,
    Candidate,
    Job,
    Application,
    SavedJob,
    ApplicationTimeline,
    AdminAuditLog
)

from .serializers import (
    JobSerializer,
    RegisterSerializer,
    CandidateSerializer,
    EmployerSerializer,
    ResumeSerializer,
    UserSerializer,
    ApplicationSerializer,
    SavedJobSerializer,
    ApplicationTimelineSerializer,
    AdminAuditLogSerializer
    
)

from .permissions import (
    IsAdmin,
    IsEmployer,
    IsCandidate
)


class TestAPIView(APIView):

    def get(self, request):
        return Response({
            "message": "Hello DRF"
        })


class SignupAPIView(APIView):

    def post(self, request):

        serializer = RegisterSerializer(
            data=request.data
        )

        if serializer.is_valid():

            serializer.save()

            return Response(
                {
                    "message": "User created successfully"
                },
                status=status.HTTP_201_CREATED
            )

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )


class LogoutAPIView(APIView):

    permission_classes = [IsAuthenticated]

    def post(self, request):

        try:
            refresh_token = request.data["refresh"]

            token = RefreshToken(refresh_token)

            token.blacklist()

            return Response({
                "message": "Logout successful"
            })

        except Exception as e:

            return Response(
                {
                    "error": str(e)
                },
                status=status.HTTP_400_BAD_REQUEST
            )




class JobCreateAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    def post(self, request):

        employer = Employer.objects.get(
            user=request.user
        )

        serializer = JobSerializer(
            data=request.data
        )

        if serializer.is_valid():

            serializer.save(
                employer=employer
            )

            return Response(
                serializer.data,
                status=status.HTTP_201_CREATED
            )

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )
class JobUpdateAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    def put(self, request, job_id):

        employer = Employer.objects.get(
            user=request.user
        )

        job = Job.objects.get(
            id=job_id
        )

        if job.employer != employer:

            return Response(
                {
                    "error": "Permission denied"
                },
                status=status.HTTP_403_FORBIDDEN
            )

        serializer = JobSerializer(
            job,
            data=request.data,
            partial=True
        )

        if serializer.is_valid():

            serializer.save()

            return Response(serializer.data)

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )
class JobStatusAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    def put(self, request, job_id):

        employer = Employer.objects.get(
            user=request.user
        )

        job = Job.objects.get(
            id=job_id
        )

        if job.employer != employer:

            return Response(
                {
                    "error": "Permission denied"
                },
                status=status.HTTP_403_FORBIDDEN
            )

        job.is_active = not job.is_active
        job.save()

        return Response(
            {
                "message": "Job status updated successfully",
                "is_active": job.is_active
            }
        )
class CandidateAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsCandidate
    ]

    def get(self, request):

        return Response({
            "message": "Candidate Dashboard"
        })


class AdminAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsAdmin
    ]

    def get(self, request):

        return Response({
            "message": "Admin Dashboard"
        })


class ApplyJobAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsCandidate
    ]

    def post(self, request, job_id):

        job = Job.objects.get(
            id=job_id
        )

        if not job.is_active:

            return Response(
                {
                    "error": "This job is no longer active."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        candidate = Candidate.objects.get(
            user=request.user
        )

        if Application.objects.filter(
            candidate=candidate,
            job=job
        ).exists():

            return Response(
                {
                    "error": "You have already applied for this job."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        application = Application.objects.create(
            candidate=candidate,
            job=job,
            resume_snapshot=candidate.resume
        )

        ApplicationTimeline.objects.create(
            application=application,
            status="Applied"
        )

        return Response(
            {
                "message": "Application submitted successfully."
            },
            status=status.HTTP_201_CREATED
        )

        Application.objects.create(
            candidate=candidate,
            job=job,
            resume_snapshot=candidate.resume
        )

        return Response(
            {
                "message": "Applied successfully"
            },
            status=status.HTTP_201_CREATED
        )


class CandidateProfileAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsCandidate
    ]

    # Create Profile
    def post(self, request):

        if Candidate.objects.filter(
            user=request.user
        ).exists():

            return Response(
                {
                    "error": "Profile already exists."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        serializer = CandidateSerializer(
            data=request.data
        )

        if serializer.is_valid():

            serializer.save(
                user=request.user
            )

            return Response(
                serializer.data,
                status=status.HTTP_201_CREATED
            )

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )

    # View Profile
    def get(self, request):

        candidate = Candidate.objects.get(
            user=request.user
        )

        serializer = CandidateSerializer(
            candidate
        )

        return Response(serializer.data)

    # Update Profile
    def put(self, request):

        candidate = Candidate.objects.get(
            user=request.user
        )

        serializer = CandidateSerializer(
            candidate,
            data=request.data,
            partial=True
        )

        if serializer.is_valid():

            serializer.save()

            return Response(
                serializer.data
            )

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )

    # Soft Delete
    def delete(self, request):

        candidate = Candidate.objects.get(
            user=request.user
        )

        candidate.is_active = False
        candidate.save()

        return Response(
            {
                "message": "Profile deactivated successfully."
            }
        )


class EmployerProfileAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    # Create Profile
    def post(self, request):

        if Employer.objects.filter(
            user=request.user
        ).exists():

            return Response(
                {
                    "error": "Profile already exists."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        serializer = EmployerSerializer(
            data=request.data
        )

        if serializer.is_valid():

            serializer.save(
                user=request.user
            )

            return Response(
                serializer.data,
                status=status.HTTP_201_CREATED
            )

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )

    # View Profile
    def get(self, request):

        employer = Employer.objects.get(
            user=request.user
        )

        serializer = EmployerSerializer(
            employer
        )

        return Response(serializer.data)

    # Update Profile
    def put(self, request):

        employer = Employer.objects.get(
            user=request.user
        )

        serializer = EmployerSerializer(
            employer,
            data=request.data,
            partial=True
        )

        if serializer.is_valid():

            serializer.save()

            return Response(
                serializer.data
            )

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )

    # Soft Delete
    def delete(self, request):

        employer = Employer.objects.get(
            user=request.user
        )

        employer.is_active = False
        employer.save()

        return Response(
            {
                "message": "Profile deactivated successfully."
            }
        )
class ResumeUploadAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsCandidate
    ]

    def post(self, request):

        candidate = Candidate.objects.get(
            user=request.user
        )

        serializer = ResumeSerializer(
            candidate,
            data=request.data,
            partial=True
        )

        if serializer.is_valid():

            serializer.save()

            return Response(
                {
                    "message": "Resume uploaded successfully",
                    "resume": candidate.resume.url
                }
            )

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )
class JobPagination(PageNumberPagination):

    page_size = 5


class JobListAPIView(ListAPIView):

    permission_classes = [IsAuthenticated]

    queryset = Job.objects.select_related(
        'employer'
    ).filter(
        is_active=True
    )

    serializer_class = JobSerializer

    pagination_class = JobPagination

    filter_backends = [
        DjangoFilterBackend,
        SearchFilter
    ]

    filterset_fields = {
        'experience': ['exact', 'gte', 'lte'],
        'salary': ['exact', 'gte', 'lte'],
        'location': ['exact'],
        'job_type': ['exact']
    }

    search_fields = [
        'title',
        'skills',
        'description'
    ]
    


class LatestJobAPIView(ListAPIView):

    permission_classes = [IsAuthenticated]

    serializer_class = JobSerializer

    pagination_class = JobPagination

    queryset = Job.objects.select_related(
        'employer'
    ).filter(
        is_active=True
    ).order_by(
        '-created_at'
    )[:10]
    
class FeaturedJobAPIView(ListAPIView):

    permission_classes = [IsAuthenticated]

    queryset = Job.objects.select_related(
        'employer'
    ).filter(
        featured=True,
        is_active=True
    )

    serializer_class = JobSerializer

    pagination_class = JobPagination


class UserListAPIView(ListAPIView):

    permission_classes = [IsAuthenticated]

    queryset = User.objects.all()

    serializer_class = UserSerializer

    filter_backends = [
        DjangoFilterBackend
    ]

    filterset_fields = [
        'role'
    ]
class ApplicationHistoryAPIView(ListAPIView):

    permission_classes = [
        IsAuthenticated,
        IsCandidate
    ]

    serializer_class = ApplicationSerializer

    def get_queryset(self):

        candidate = Candidate.objects.get(
            user=self.request.user
        )

        return Application.objects.select_related(
            'job',
            'candidate'
        ).filter(
            candidate=candidate
        ).order_by(
            '-applied_at'
        )
class ApplicationStatusAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    def put(self, request, application_id):

        employer = Employer.objects.get(
            user=request.user
        )

        application = Application.objects.get(
            id=application_id
        )

        if application.job.employer != employer:

            return Response(
                {
                    "error": "Permission denied."
                },
                status=status.HTTP_403_FORBIDDEN
            )

        new_status = request.data.get(
            "status",
            ""
        ).strip()

        valid_transitions = {

            "Applied": [
                "Under Review",
                "Rejected"
            ],

            "Under Review": [
                "Shortlisted",
                "Rejected"
            ],

            "Shortlisted": [
                "Interview Scheduled",
                "Rejected"
            ],

            "Interview Scheduled": [
                "Selected",
                "Rejected"
            ],

            "Selected": [],

            "Rejected": []
        }

        if new_status not in valid_transitions[
            application.status
        ]:

            return Response(
                {
                    "error": "Invalid status transition."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        application.status = new_status

        application.save()

        # Save timeline
        ApplicationTimeline.objects.create(
            application=application,
            status=new_status
        )

        return Response(
            {
                "message": "Application status updated successfully.",
                "status": application.status,
                "updated_at": application.updated_at
            }
        )
class EmployerJobListAPIView(ListAPIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    serializer_class = JobSerializer

    def get_queryset(self):

        employer = Employer.objects.get(
            user=self.request.user
        )

        return Job.objects.select_related(
            'employer'
        ).filter(
            employer=employer
        )

class ApplicantListAPIView(ListAPIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    serializer_class = ApplicationSerializer

    filter_backends = [
        SearchFilter
    ]

    search_fields = [
        "candidate__full_name"
    ]

    def get_queryset(self):

        employer = Employer.objects.get(
            user=self.request.user
        )

        job = Job.objects.get(
            id=self.kwargs["job_id"]
        )

        # Check job ownership
        if job.employer != employer:

            raise PermissionDenied(
                "Permission denied."
            )

        queryset = Application.objects.select_related(
            "candidate",
            "job"
        ).filter(
            job=job
        )

        print("Before filter:", queryset.count())

        status = self.request.query_params.get(
            "status"
        )

        print("Status received:", status)

        if status:

            queryset = queryset.filter(
                status=status
            )

        print("After filter:", queryset.count())

        return queryset
class EmployerDashboardAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    def get(self, request):

        employer = Employer.objects.get(
            user=request.user
        )

        jobs = Job.objects.filter(
            employer=employer
        )

        total_jobs = jobs.count()

        active_jobs = jobs.filter(
            is_active=True
        ).count()

        total_applications = Application.objects.filter(
            job__employer=employer
        ).count()

        shortlisted_candidates = Application.objects.filter(
            job__employer=employer,
            status="Shortlisted"
        ).count()

        return Response(
            {
                "total_jobs": total_jobs,
                "active_jobs": active_jobs,
                "total_applications": total_applications,
                "shortlisted_candidates": shortlisted_candidates
            }
        )
class SaveJobAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsCandidate
    ]

    def post(self, request, job_id):

        candidate = Candidate.objects.get(
            user=request.user
        )

        job = Job.objects.get(
            id=job_id
        )

        if SavedJob.objects.filter(
            candidate=candidate,
            job=job
        ).exists():

            return Response(
                {
                    "error": "Job already saved."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        SavedJob.objects.create(
            candidate=candidate,
            job=job
        )

        return Response(
            {
                "message": "Job saved successfully."
            },
            status=status.HTTP_201_CREATED
        )
class SavedJobListAPIView(ListAPIView):

    permission_classes = [
        IsAuthenticated,
        IsCandidate
    ]

    serializer_class = SavedJobSerializer

    def get_queryset(self):

        candidate = Candidate.objects.get(
            user=self.request.user
        )

        return SavedJob.objects.select_related(
            "job",
            "job__employer"
        ).filter(
            candidate=candidate
        ).order_by(
            "-saved_at"
        )
class RecommendedJobAPIView(ListAPIView):

    permission_classes = [
        IsAuthenticated,
        IsCandidate
    ]

    serializer_class = JobSerializer

    def get_queryset(self):

        candidate = Candidate.objects.get(
            user=self.request.user
        )

        skills = [
            skill.strip().lower()
            for skill in candidate.skills.split(",")
        ]

        queryset = Job.objects.select_related(
            "employer"
        ).filter(
            is_active=True
        )

        recommended_jobs = []

        for job in queryset:

            job_skills = [
                skill.strip().lower()
                for skill in job.skills.split(",")
            ]

            if any(
                skill in job_skills
                for skill in skills
            ):
                recommended_jobs.append(
                    job.id
                )

        return Job.objects.filter(
            id__in=recommended_jobs
        )
class ApplicationTimelineAPIView(ListAPIView):

    permission_classes = [
        IsAuthenticated,
        IsCandidate
    ]

    serializer_class = ApplicationTimelineSerializer

    def get_queryset(self):

        candidate = Candidate.objects.get(
            user=self.request.user
        )

        application = Application.objects.get(
            id=self.kwargs["application_id"],
            candidate=candidate
        )

        return ApplicationTimeline.objects.filter(
            application=application
        ).order_by(
            "changed_at"
        )
class ApproveEmployerAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsAdmin
    ]

    def put(self, request, employer_id):

        employer = Employer.objects.get(
            id=employer_id
        )

        employer.verified = True

        employer.save()

        # Insert audit log
        AdminAuditLog.objects.create(
            admin=request.user,
            action=f"Approved employer {employer.company_name}"
        )

        return Response(
            {
                "message": "Employer approved successfully.",
                "verified": employer.verified
            },
            status=status.HTTP_200_OK
        )
class BlockUserAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsAdmin
    ]

    def put(self, request, user_id):

        user = User.objects.get(
            id=user_id
        )

        user.is_active = False

        user.save()
        AdminAuditLog.objects.create(
        admin=request.user,
        action=f"Blocked user {user.username}"
        )

        return Response(
            {
                "message": "User blocked successfully.",
                "is_active": user.is_active
            },
            status=status.HTTP_200_OK
        )
class UnblockUserAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsAdmin
    ]

    def put(self, request, user_id):

        user = User.objects.get(
            id=user_id
        )

        user.is_active = True

        user.save()

        AdminAuditLog.objects.create(
            admin=request.user,
            action=f"Unblocked user '{user.username}'"
        )

        return Response(
            {
                "message": "User unblocked successfully.",
                "is_active": user.is_active
            },
            status=status.HTTP_200_OK
        )
class AdminJobListAPIView(ListAPIView):

    permission_classes = [
        IsAuthenticated,
        IsAdmin
    ]

    serializer_class = JobSerializer

    queryset = Job.objects.select_related(
        "employer"
    ).all()
class DeleteJobAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsAdmin
    ]

    def delete(self, request, job_id):

        job = Job.objects.get(
            id=job_id
        )

        job_title = job.title

        job.delete()

        AdminAuditLog.objects.create(
            admin=request.user,
            action=f"Deleted job '{job_title}'"
        )

        return Response(
            {
                "message": "Job deleted successfully."
            },
            status=status.HTTP_200_OK
        )


class AdminJobUpdateAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsAdmin
    ]

    def put(self, request, job_id):

        job = Job.objects.get(
            id=job_id
        )

        serializer = JobSerializer(
            job,
            data=request.data,
            partial=True
        )

        if serializer.is_valid():

            serializer.save()

            AdminAuditLog.objects.create(
                admin=request.user,
                action=f"Updated job '{job.title}'"
            )

            return Response(
                {
                    "message": "Job updated successfully.",
                    "data": serializer.data
                },
                status=status.HTTP_200_OK
            )

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )


class AdminPlatformStatsAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsAdmin
    ]

    def get(self, request):

        total_users = User.objects.count()

        total_employers = Employer.objects.count()

        total_candidates = Candidate.objects.count()

        total_jobs = Job.objects.count()

        active_jobs = Job.objects.filter(
            is_active=True
        ).count()

        inactive_jobs = Job.objects.filter(
            is_active=False
        ).count()

        total_applications = Application.objects.count()

        return Response(
            {
                "total_users": total_users,
                "total_employers": total_employers,
                "total_candidates": total_candidates,
                "total_jobs": total_jobs,
                "active_jobs": active_jobs,
                "inactive_jobs": inactive_jobs,
                "total_applications": total_applications
            },
            status=status.HTTP_200_OK
        )


class AdminAuditLogAPIView(ListAPIView):

    permission_classes = [
        IsAuthenticated,
        IsAdmin
    ]

    serializer_class = AdminAuditLogSerializer

    queryset = AdminAuditLog.objects.select_related(
        "admin"
    ).order_by(
        "-created_at"
    )
class ResumeParserAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsCandidate
    ]

    def get(self, request):

        candidate = Candidate.objects.get(
            user=request.user
        )

        if not candidate.resume:

            return Response(
                {
                    "error": "Resume not uploaded."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        file_path = candidate.resume.path

        extracted_text = ""

        if file_path.endswith(".pdf"):

            with pdfplumber.open(file_path) as pdf:

                for page in pdf.pages:

                    text = page.extract_text()

                    if text:

                        extracted_text += text + "\n"

        elif file_path.endswith(".docx"):

            document = docx.Document(file_path)

            for paragraph in document.paragraphs:

                extracted_text += paragraph.text + "\n"

        else:

            return Response(
                {
                    "error": "Unsupported file format."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        cleaned_text = " ".join(
            extracted_text.split()
        )

        return Response(
            {
                "resume_text": cleaned_text
            }
        )
class ParsedResumeAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsCandidate
    ]

    def get(self, request):

        candidate = Candidate.objects.get(
            user=request.user
        )

        if not candidate.resume:

            return Response(
                {
                    "error": "Resume not uploaded."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        file_path = candidate.resume.path

        extracted_text = ""

        if file_path.endswith(".pdf"):

            with pdfplumber.open(file_path) as pdf:

                for page in pdf.pages:

                    text = page.extract_text()

                    if text:

                        extracted_text += text + "\n"

        elif file_path.endswith(".docx"):

            document = docx.Document(file_path)

            for paragraph in document.paragraphs:

                extracted_text += paragraph.text + "\n"

        else:

            return Response(
                {
                    "error": "Unsupported file format."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        cleaned_text = " ".join(
            extracted_text.split()
        )

        tokens = cleaned_text.split()

        name = candidate.full_name

        email = re.search(
            r'[\w\.-]+@[\w\.-]+\.\w+',
            cleaned_text
        )

        email = email.group() if email else ""

        phone = re.search(
            r'\b\d{10}\b',
            cleaned_text
        )

        phone = phone.group() if phone else ""

        predefined_skills = [
            "Python",
            "Java",
            "C",
            "C++",
            "PHP",
            "Django",
            "Flutter",
            "HTML",
            "CSS",
            "Bootstrap",
            "JavaScript",
            "SQL",
            "MySQL",
            "Git",
            "GitHub",
            "NumPy",
            "Pandas",
            "Scikit-Learn",
            "REST API"
        ]

        extracted_skills = []

        for skill in predefined_skills:

            if skill.lower() in cleaned_text.lower():

                extracted_skills.append(
                    skill
                )

        education = ""

        if "MCA" in cleaned_text:

            education = "Master of Computer Applications"

        elif "BCA" in cleaned_text:

            education = "Bachelor of Computer Applications"

        elif "B.Tech" in cleaned_text:

            education = "Bachelor of Technology"

        elif "B.A" in cleaned_text:

            education = "Bachelor of Arts"

        experience = re.search(
            r'(\d+)\s+Years?',
            cleaned_text,
            re.IGNORECASE
        )

        if experience:

            experience = experience.group()

        else:

            experience = "Fresher"

        return Response(
            {
                "name": name,
                "email": email,
                "phone": phone,
                "skills": extracted_skills,
                "education": education,
                "experience": experience
            },
            status=status.HTTP_200_OK
        )
class ATSScoreAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    def get(self, request, application_id):

        application = Application.objects.get(
            id=application_id
        )

        employer = Employer.objects.get(
            user=request.user
        )

        if application.job.employer != employer:

            return Response(
                {
                    "error": "Permission denied."
                },
                status=status.HTTP_403_FORBIDDEN
            )

        candidate = application.candidate

        job = application.job

        job_skills = [
            skill.strip().lower()
            for skill in job.skills.split(",")
        ]

        candidate_skills = [
            skill.strip().lower()
            for skill in candidate.skills.split(",")
        ]

        matched_skills = 0

        for skill in job_skills:

            if skill in candidate_skills:

                matched_skills += 1

        if len(job_skills) > 0:

            skill_score = (
                matched_skills / len(job_skills)
            ) * 60

        else:

            skill_score = 0

        if candidate.experience >= job.experience:

            experience_score = 25

        else:

            experience_score = (
                candidate.experience /
                job.experience
            ) * 25 if job.experience > 0 else 0

        if "MCA" in candidate.education.upper():

            education_score = 15

        else:

            education_score = 0

        total_score = round(

            skill_score +
            experience_score +
            education_score,

            2

        )

        application.ats_score = total_score

        application.save()

        return Response(
    {
        "candidate": candidate.full_name,
        "job": job.title,
        "matched_skills": matched_skills,
        "total_job_skills": len(job_skills),
        "skill_score": skill_score,
        "experience_score": experience_score,
        "education_score": education_score,
        "ats_score": total_score
    }
)
class RankedCandidateAPIView(ListAPIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    serializer_class = ApplicationSerializer

    def get_queryset(self):

        employer = Employer.objects.get(
            user=self.request.user
        )

        job = Job.objects.get(
            id=self.kwargs["job_id"]
        )

        if job.employer != employer:

            raise PermissionDenied(
                "Permission denied."
            )

        return Application.objects.select_related(
            "candidate",
            "job"
        ).filter(
            job=job
        ).order_by(
            "-ats_score"
        )
class AutoHiringAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    def put(self, request, application_id):

        employer = Employer.objects.get(
            user=request.user
        )

        application = Application.objects.get(
            id=application_id
        )

        if application.job.employer != employer:

            return Response(
                {
                    "error": "Permission denied."
                },
                status=status.HTTP_403_FORBIDDEN
            )

        score = application.ats_score

        if score >= 80:

            application.status = "Shortlisted"

        elif score >= 60:

            application.status = "Under Review"

        else:

            application.status = "Rejected"

        application.save()

        return Response(
            {
                "candidate": application.candidate.full_name,
                "job": application.job.title,
                "ats_score": score,
                "status": application.status
            },
            status=status.HTTP_200_OK
        )
class EmployerOverrideAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    def put(self, request, application_id):

        employer = Employer.objects.get(
            user=request.user
        )

        application = Application.objects.get(
            id=application_id
        )

        if application.job.employer != employer:

            return Response(
                {
                    "error": "Permission denied."
                },
                status=status.HTTP_403_FORBIDDEN
            )

        new_status = request.data.get("status")

        valid_status = [

            "Applied",
            "Under Review",
            "Shortlisted",
            "Interview Scheduled",
            "Selected",
            "Rejected"

        ]

        if new_status not in valid_status:

            return Response(
                {
                    "error": "Invalid status."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        application.status = new_status

        application.save()

        return Response(
            {
                "message": "Application status updated by employer.",
                "candidate": application.candidate.full_name,
                "job": application.job.title,
                "status": application.status
            },
            status=status.HTTP_200_OK
        )
class AutoNotificationAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    def get(self, request, application_id):

        employer = Employer.objects.get(
            user=request.user
        )

        application = Application.objects.get(
            id=application_id
        )

        if application.job.employer != employer:

            return Response(
                {
                    "error": "Permission denied."
                },
                status=status.HTTP_403_FORBIDDEN
            )

        if application.status == "Shortlisted":

            message = (
                f"Congratulations {application.candidate.full_name}! "
                f"You have been shortlisted for the job "
                f"'{application.job.title}'."
            )

        elif application.status == "Rejected":

            message = (
                f"Dear {application.candidate.full_name}, "
                f"thank you for applying for "
                f"'{application.job.title}'. "
                f"Unfortunately, your application was not selected."
            )

        elif application.status == "Interview Scheduled":

            message = (
                f"Hello {application.candidate.full_name}, "
                f"your interview has been scheduled for "
                f"'{application.job.title}'."
            )

        elif application.status == "Selected":

            message = (
                f"Congratulations {application.candidate.full_name}! "
                f"You have been selected for "
                f"'{application.job.title}'."
            )

        else:

            message = (
                f"Your application for "
                f"'{application.job.title}' "
                f"is currently '{application.status}'."
            )

        return Response(
            {
                "candidate": application.candidate.full_name,
                "job": application.job.title,
                "status": application.status,
                "notification": message
            },
            status=status.HTTP_200_OK
        )
class BatchAutoHiringAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    def put(self, request, job_id):

        employer = Employer.objects.get(
            user=request.user
        )

        job = Job.objects.get(
            id=job_id
        )

        if job.employer != employer:

            return Response(
                {
                    "error": "Permission denied."
                },
                status=status.HTTP_403_FORBIDDEN
            )

        applications = Application.objects.filter(
            job=job
        )

        updated = 0

        for application in applications:

            score = application.ats_score

            if score >= 80:

                application.status = "Shortlisted"

            elif score >= 60:

                application.status = "Under Review"

            else:

                application.status = "Rejected"

            application.save()

            updated += 1

        return Response(
            {
                "message": "Batch processing completed successfully.",
                "job": job.title,
                "applications_processed": updated
            },
            status=status.HTTP_200_OK
        )
class SendEmailAPIView(APIView):

    permission_classes = [
        IsAuthenticated,
        IsEmployer
    ]

    def post(self, request, application_id):

        employer = Employer.objects.get(
            user=request.user
        )

        application = Application.objects.get(
            id=application_id
        )

        if application.job.employer != employer:

            return Response(
                {
                    "error": "Permission denied."
                },
                status=status.HTTP_403_FORBIDDEN
            )

        if application.status == "Shortlisted":

            subject = "Application Shortlisted"

            message = (
                f"Congratulations {application.candidate.full_name},\n\n"
                f"You have been shortlisted for the position "
                f"{application.job.title}."
            )

        elif application.status == "Rejected":

            subject = "Application Update"

            message = (
                f"Dear {application.candidate.full_name},\n\n"
                f"Thank you for applying for "
                f"{application.job.title}.\n"
                f"Unfortunately, your application was not selected."
            )

        elif application.status == "Selected":

            subject = "Congratulations!"

            message = (
                f"Congratulations {application.candidate.full_name},\n\n"
                f"You have been selected for "
                f"{application.job.title}."
            )

        else:

            subject = "Application Status"

            message = (
                f"Hello {application.candidate.full_name},\n\n"
                f"Your application status is "
                f"{application.status}."
            )

        send_mail(

            subject,

            message,

            settings.EMAIL_HOST_USER,

            [application.candidate.user.email],

            fail_silently=False

        )

        return Response(
            {
                "message": "Email sent successfully."
            },
            status=status.HTTP_200_OK
        )
        